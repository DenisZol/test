#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import re
import sys
import shutil
import pathlib
from datetime import datetime, timedelta

import pdfplumber
from docx import Document

CONSTANT_TEXT = "XXX"
TEMPLATE_DOCX = "Письмо на Банк Благотворит.docx"
TARGET_ROOT   = pathlib.Path.cwd()

RE_DATE      = re.compile(r'(\d{1,2})/(\d{1,2})/(\d{4})')
RE_INVNO     = re.compile(r'Invoice\s*No\.?\s*0*(\d{3,})', re.I)
RE_INVNO_FBK = re.compile(r'\b000\d{5}\b')
RE_TOTALUSD  = re.compile(r'Total\s+amount:?\s*USD\s+([\d\s,]+(?:\.\d{2})?)', re.I)
RE_ANYUSD    = re.compile(r'USD\s+([\d\s,]+(?:\.\d{2})?)', re.I)

UA_MONTHS = {
    1:"січня", 2:"лютого", 3:"березня", 4:"квітня",
    5:"травня", 6:"червня", 7:"липня", 8:"серпня",
    9:"вересня", 10:"жовтня", 11:"листопада", 12:"грудня"
}
def ua_date(dt: datetime) -> str:
    return f"{dt.day} {UA_MONTHS[dt.month]} {dt.year}"

def fill_docx(template: pathlib.Path, dest: pathlib.Path,
              plus2_dt: datetime, amount_full: str):
    doc = Document(template)
    mapping = {
        "{{DATE}}":        ua_date(plus2_dt),
        "{{DATE + 1}}":    ua_date(plus2_dt),
        "{{FULL_AMOUNT}}": amount_full,
    }
    for p in doc.paragraphs:
        for k, v in mapping.items():
            if k in p.text:
                for r in p.runs:
                    r.text = r.text.replace(k, v)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for k, v in mapping.items():
                        if k in para.text:
                            for r in para.runs:
                                r.text = r.text.replace(k, v)
    doc.save(dest)

def extract_from_pdf(pdf_path: pathlib.Path):
    with pdfplumber.open(pdf_path) as pdf:
        text = "\n".join(p.extract_text() or "" for p in pdf.pages)
    m_date = RE_DATE.search(text)
    if not m_date:
        raise ValueError("Дата (M/D/YYYY) не найдена")
    month, day, year = map(int, m_date.groups())
    invoice_dt = datetime(year, month, day)

    m_num = RE_INVNO.search(text) or RE_INVNO_FBK.search(text)
    if not m_num:
        raise ValueError("Номер инвойса не найден")
    invoice_num = m_num.group(1) if m_num.re is RE_INVNO else m_num.group(0).lstrip('0')

    m_amt = RE_TOTALUSD.search(text) or RE_ANYUSD.search(text)
    if not m_amt:
        raise ValueError("USD amount не найдена")
    raw_amt = m_amt.group(1)
    norm = raw_amt.replace(' ', '').replace(',', '')
    amount_float = float(norm)
    amount_full  = f"{amount_float:,.2f}".replace(',', ' ').replace('.', ',')
    amount_int   = int(amount_float)

    return invoice_dt, invoice_num, amount_full, amount_int

def process_pdf(pdf_path: pathlib.Path):
    print(f"\n=== Обрабатываю: {pdf_path.name} ===")
    invoice_dt, inv_num, amount_full, amount_int = extract_from_pdf(pdf_path)
    plus2_dt = invoice_dt + timedelta(days=2)

    folder = f"{plus2_dt.strftime('%y-%m')} {CONSTANT_TEXT} {amount_int} №{inv_num}"
    folder = "".join(c for c in folder if c not in r'<>:"/\\|?*').strip(" .")
    target_dir = TARGET_ROOT / folder
    target_dir.mkdir(parents=True, exist_ok=True)
    print(f"Создана (или уже была): {target_dir}")

    template = TARGET_ROOT / TEMPLATE_DOCX
    if template.exists():
        out_doc = target_dir / f"Лист_в_банк_{inv_num}.docx"
        fill_docx(template, out_doc, plus2_dt, amount_full)
        print(f"Створено Word: {out_doc}")
    else:
        print(f"⚠️  Шаблон {TEMPLATE_DOCX} не найден — пропускаю генерацию Word.")

    dest_pdf = target_dir / pdf_path.name
    if not dest_pdf.exists():
        shutil.move(str(pdf_path), dest_pdf)
        print(f"PDF перемещён в: {dest_pdf}")
    else:
        print("PDF уже лежит в целевой папке — не перемещаю.")

def main():
    pdf_files = sorted(TARGET_ROOT.glob("Invoice*.pdf"))
    if not pdf_files:
        print("Нет файлов, начинающихся с 'Invoice' и заканчивающихся на .pdf")
        return
    for pdf in pdf_files:
        try:
            process_pdf(pdf)
        except Exception as e:
            print(f"❌ Ошибка при обработке {pdf.name}: {e}")

if __name__ == "__main__":
    main()
